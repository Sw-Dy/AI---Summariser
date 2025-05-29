<<<<<<< HEAD
import re
import json
import logging
import nltk
from transformers import pipeline
from textwrap import wrap
from langdetect import detect
from nltk.corpus import stopwords
from textblob import TextBlob
from deep_translator import GoogleTranslator
from nltk.tokenize import sent_tokenize, word_tokenize
from rake_nltk import Rake
import spacy
import PyPDF2
import pytesseract
from PIL import Image
import whisper
from fpdf import FPDF
from docx import Document
import textstat
import easyocr 
import speech_recognition as sr
import torch
import openai
import os
from pydub import AudioSegment
import librosa
import torch
from transformers import Wav2Vec2ForCTC, Wav2Vec2Tokenizer
import cv2
from transformers import VisionEncoderDecoderModel, ViTImageProcessor, AutoTokenizer


# Initialize necessary components
logging.basicConfig(filename="summarizer.log", level=logging.ERROR, format="%(asctime)s - %(levelname)s - %(message)s")
models = {
    "t5-small": "t5-small",
    "bart-large": "facebook/bart-large-cnn",
    "distilbart": "sshleifer/distilbart-cnn-12-6",
    "pegasus": "google/pegasus-xsum",
    "t5-3b": "t5-3b"
}
stopwords.words("english")  # Download if not already done
nlp = spacy.load("en_core_web_sm")

# Preprocessing Functions
def preprocess_text(text, preserve_case=True, keep_punctuation=True):
    """Enhanced preprocessing with options to preserve case and punctuation for better ROUGE scores"""
    # Initial cleaning
    text = text.strip()
    
    if not preserve_case:
        text = text.lower()
    
    if keep_punctuation:
        # Only remove excessive punctuation and normalize spacing
        text = re.sub(r'([.,!?;:]){2,}', r'\1', text)  # Replace multiple punctuation with single
        text = re.sub(r'\s+', ' ', text).strip()  # Normalize spacing
    else:
        # Remove all punctuation
        text = re.sub(r'[^a-zA-Z0-9\s]', '', text)
        text = re.sub(r'\s+', ' ', text).strip()
    
    # Tokenize and filter stopwords if needed
    words = word_tokenize(text)
    
    # Only filter stopwords for certain summarization approaches
    # Keeping stopwords often helps with ROUGE scores
    filtered_words = [word for word in words if word.lower() not in stopwords.words("english") or len(word) <= 2 or not word.isalpha()]
    
    return " ".join(filtered_words)

def chunk_text(text, chunk_size=500, overlap=50, respect_sentences=True):
    """Enhanced text chunking with overlap and sentence boundary respect for better coherence"""
    if not text:
        return []
        
    if not respect_sentences:
        # Simple chunking with overlap
        chunks = []
        for i in range(0, len(text), chunk_size - overlap):
            chunks.append(text[i:i + chunk_size])
        return chunks
    
    # Sentence-aware chunking
    sentences = sent_tokenize(text)
    chunks = []
    current_chunk = ""
    
    for sentence in sentences:
        # If adding this sentence would exceed chunk size and we already have content
        if len(current_chunk) + len(sentence) > chunk_size and current_chunk:
            chunks.append(current_chunk.strip())
            # Keep some overlap by including the last sentence in the next chunk
            last_sentences = sent_tokenize(current_chunk)[-1] if sent_tokenize(current_chunk) else ""
            current_chunk = last_sentences + " " + sentence if last_sentences else sentence
        else:
            current_chunk += " " + sentence if current_chunk else sentence
    
    # Add the last chunk if it's not empty
    if current_chunk.strip():
        chunks.append(current_chunk.strip())
        
    return chunks

# Summarization Functions
def summarize_text(text, model_choice, max_words=None, optimize_for_rouge=True):
    """Enhanced summarization function with ROUGE optimization"""
    try:
        # Get the maximum word limit from user input if not provided
        if max_words is None:
            max_words = int(input("Enter maximum word limit for the summary (recommended: 20-30% of original length): "))
            
        # For t5-3b, warn about memory requirements and set smaller chunk size
        if model_choice == "t5-3b":
            logging.info("Using t5-3b model - this requires significant memory and may be slower")
            # Reduce memory usage by processing smaller chunks for large models
            chunk_size = 512  # Smaller chunks for t5-3b
            overlap = 50
        else:
            chunk_size = 1000
            overlap = 100
        
        # Calculate original text length for better proportional summarization
        original_word_count = len(text.split())
        
        # Adjust min_length and max_length based on original text length
        # Ensure min_length is always less than max_length
        max_length = min(max_words, int(original_word_count * 0.4))  # At most 40% of original
        min_length = min(max(5, int(original_word_count * 0.05)), max_length - 1)  # At least 5% of original but less than max_length
        
        # Initialize the summarization model
        model_name = models.get(model_choice, "facebook/bart-large-cnn")
        
        # Special handling for large models like t5-3b
        if model_choice == "t5-3b":
            try:
                from transformers import AutoModelForSeq2SeqLM, AutoTokenizer
                # Load with device_map="auto" for large models to distribute across available devices
                tokenizer = AutoTokenizer.from_pretrained(model_name)
                
                # Try to load with optimal settings first
                try:
                    t5_model = AutoModelForSeq2SeqLM.from_pretrained(
                        model_name,
                        device_map="auto",  # Automatically distribute model across available devices
                        torch_dtype=torch.float16,  # Use half precision to reduce memory usage
                        low_cpu_mem_usage=True
                    )
                except Exception as e:
                    logging.warning(f"Could not load t5-3b with optimal settings: {e}. Trying with reduced settings.")
                    # Fallback to more conservative settings
                    t5_model = AutoModelForSeq2SeqLM.from_pretrained(
                        model_name,
                        torch_dtype=torch.float16,
                        low_cpu_mem_usage=True
                    )
                
                # Use gradient checkpointing to reduce memory usage during forward pass
                t5_model.gradient_checkpointing_enable()
                
                model = pipeline("summarization", model=t5_model, tokenizer=tokenizer, framework="pt")
                logging.info("Successfully loaded t5-3b model")
            except Exception as e:
                logging.error(f"Failed to load t5-3b model: {e}. Falling back to bart-large-cnn.")
                # Fallback to a smaller model
                model_name = "facebook/bart-large-cnn"
                model = pipeline("summarization", model=model_name, tokenizer=model_name, framework="pt", device=0 if torch.cuda.is_available() else -1)
        else:
            # Standard pipeline for smaller models
            model = pipeline("summarization", model=model_name, tokenizer=model_name, framework="pt", device=0 if torch.cuda.is_available() else -1)
        
        # Use enhanced text chunking with overlap and sentence boundaries
        text_chunks = chunk_text(text, chunk_size=chunk_size, overlap=overlap, respect_sentences=True)
        
        # Process each chunk with optimized parameters
        summaries = []
        for i, chunk in enumerate(text_chunks):
            try:
                # Adjust parameters based on chunk position for better coherence
                chunk_min_length = min_length if len(text_chunks) == 1 else max(5, min_length // len(text_chunks))
                chunk_max_length = max_length if len(text_chunks) == 1 else max_length // len(text_chunks) + 10
                
                # Use different parameters for ROUGE optimization if requested
                if model_choice == "t5-3b":
                    # Memory-efficient generation for t5-3b
                    summary_output = model(
                        chunk, 
                        max_length=chunk_max_length,
                        min_length=chunk_min_length,
                        do_sample=True,
                        top_p=0.92,       # Higher top_p for t5-3b to improve quality
                        top_k=50,
                        num_beams=2,      # Reduced beam size to save memory
                        early_stopping=True,  # Stop early when possible to save computation
                        no_repeat_ngram_size=3  # Prevent repetition
                    )
                elif optimize_for_rouge:
                    summary_output = model(
                        chunk, 
                        max_length=chunk_max_length,
                        min_length=chunk_min_length,
                        do_sample=True,  # Enable sampling for diversity
                        top_p=0.85,      # Nucleus sampling for better quality
                        top_k=50,        # Limit vocabulary for more focused summaries
                        num_beams=4      # Beam search for better quality
                    )
                else:
                    summary_output = model(
                        chunk, 
                        max_length=chunk_max_length,
                        min_length=chunk_min_length,
                        do_sample=False
                    )
                
                summaries.append(summary_output[0]['summary_text'])
            except Exception as e:
                logging.error(f"Summarization error on chunk {i}: {e}")
                # Create a fallback summary instead of using the whole chunk
                fallback = ' '.join(chunk.split()[:chunk_max_length])
                summaries.append(fallback)
        
        # Post-process and combine summaries for better coherence
        processed_summaries = []
        for i, summary in enumerate(summaries):
            # Clean up the summary
            clean_summary = summary.strip()
            
            # Ensure the summary ends with proper punctuation
            if clean_summary and clean_summary[-1] not in '.!?':
                clean_summary += '.'
                
            processed_summaries.append(clean_summary)
        
        # Join summaries with proper spacing and connectors for better flow
        if len(processed_summaries) > 1:
            # Add transition phrases between chunks for better coherence
            transitions = ["", "Furthermore, ", "Additionally, ", "Moreover, ", "In addition, "]
            final_summary = processed_summaries[0]
            
            for i, summary in enumerate(processed_summaries[1:], 1):
                transition = transitions[min(i, len(transitions)-1)]
                final_summary += " " + transition + summary
        else:
            final_summary = processed_summaries[0] if processed_summaries else ""
        
        # Ensure the summary doesn't exceed the max word limit
        final_summary_words = final_summary.split()
        if len(final_summary_words) > max_words:
            final_summary = " ".join(final_summary_words[:max_words])
            # Ensure the truncated summary ends with proper punctuation
            if not final_summary.endswith(('.', '!', '?')):
                final_summary += '.'
        
        # Calculate and log ROUGE scores if original text is available
        try:
            from rouge_score import rouge_scorer
            scorer = rouge_scorer.RougeScorer(['rouge1', 'rouge2', 'rougeL'], use_stemmer=True)
            scores = scorer.score(text, final_summary)
            logging.info(f"ROUGE Scores - R1: {scores['rouge1'].fmeasure:.4f}, R2: {scores['rouge2'].fmeasure:.4f}, RL: {scores['rougeL'].fmeasure:.4f}")
        except Exception as e:
            logging.warning(f"Could not calculate ROUGE scores: {e}")
        
        return final_summary
    except Exception as e:
        logging.error(f"Summarization error: {e}")
        return "Error in summarization. Check logs."


# Additional Features
def extract_keywords(text):
    r = Rake()
    r.extract_keywords_from_text(text)
    return r.get_ranked_phrases()

def extract_entities(text):
    doc = nlp(text)
    return [(ent.text, ent.label_) for ent in doc.ents]

def analyze_sentiment(text):
    return TextBlob(text).sentiment.polarity

def detect_bias(text):
    classifier = pipeline("text-classification", model="facebook/bart-large-mnli")
    labels = ["biased", "neutral"]
    result = classifier(text, candidate_labels=labels)
    return result

def classify_topic(text):
    classifier = pipeline("zero-shot-classification")
    labels = ["Finance", "Health", "Technology", "Education"]
    result = classifier(text, candidate_labels=labels)
    return result["labels"][0]

def compression_ratio(original, summary):
    return (1 - len(summary) / len(original)) * 100

def readability_score(text):
    return round(textstat.flesch_reading_ease(text), 2)  # Round for consistency

# Evaluation Metrics
def evaluate_rouge_scores(reference, summary):
    """Calculate ROUGE scores between reference text and summary"""
    try:
        from rouge_score import rouge_scorer
        scorer = rouge_scorer.RougeScorer(['rouge1', 'rouge2', 'rougeL'], use_stemmer=True)
        scores = scorer.score(reference, summary)
        return {
            "rouge1": scores["rouge1"].fmeasure,
            "rouge2": scores["rouge2"].fmeasure,
            "rougeL": scores["rougeL"].fmeasure
        }
    except Exception as e:
        logging.error(f"Error calculating ROUGE scores: {e}")
        return {"rouge1": 0, "rouge2": 0, "rougeL": 0}

def evaluate_similarity(reference, summary):
    """Calculate semantic similarity between reference and summary"""
    try:
        from sklearn.feature_extraction.text import TfidfVectorizer
        from sklearn.metrics.pairwise import cosine_similarity
        
        # Create TF-IDF vectors
        vectorizer = TfidfVectorizer()
        tfidf_matrix = vectorizer.fit_transform([reference, summary])
        
        # Calculate cosine similarity
        similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
        return similarity
    except Exception as e:
        logging.error(f"Error calculating similarity: {e}")
        return 0.0

def evaluate_summary(original_text, summary_text):
    """Comprehensive evaluation of summary quality"""
    metrics = {}
    
    # Basic metrics
    metrics["compression_ratio"] = compression_ratio(original_text, summary_text)
    metrics["readability"] = readability_score(summary_text)
    
    # ROUGE scores
    rouge_scores = evaluate_rouge_scores(original_text, summary_text)
    metrics.update(rouge_scores)
    
    # Semantic similarity
    metrics["similarity"] = evaluate_similarity(original_text, summary_text)
    
    return metrics


# File Handling Functions
def read_pdf(file_path):
    try:
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            return "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
    except Exception as e:
        logging.error(f"PDF reading error: {e}")
        return "Error reading PDF file."

import easyocr
import logging
import os

# Configure logging
logging.basicConfig(level=logging.WARNING, format="%(levelname)s: %(message)s")

def extract_text_from_image(image_path):
    try:
        # Initialize EasyOCR reader (English language, using GPU if available)
        reader = easyocr.Reader(['en'], gpu=True)

        # Perform OCR
        results = reader.readtext(image_path, detail=1)  # Returns text along with confidence

        
        

        extracted_text = []

        for entry in results:
            if len(entry) >= 2:  # Ensure structure contains text and confidence
                text, confidence = entry[1], entry[2]
                
                # Apply confidence threshold (e.g., 60%)
                if confidence > 0.6:
                    extracted_text.append(text)

        final_text = "\n".join(extracted_text)  # Join text with new lines

        return final_text if final_text.strip() else "No text detected."

    except Exception as e:
        return f"Error extracting text from image: {str(e)}"

def transcribe_audio(file_path: str) -> str:
    """
    Transcribe an audio file using Wav2Vec2 model for speech recognition.

    :param file_path: Path to the audio file (WAV or other formats like MP3).
    :return: Transcribed text as a string or an error message.
    """
    try:
        if not os.path.exists(file_path):
            return "Error: Audio file not found. Please check the file path."

        # Load pre-trained model and tokenizer
        tokenizer = Wav2Vec2Tokenizer.from_pretrained("facebook/wav2vec2-base-960h")
        model = Wav2Vec2ForCTC.from_pretrained("facebook/wav2vec2-base-960h")
        
        # Load audio
        sample_rate = 16000
        speech, rate = librosa.load(file_path, sr=sample_rate)
        
        # Define chunk size (30 seconds)
        chunk_duration = 30
        samples_per_chunk = chunk_duration * sample_rate

        # Process and transcribe each chunk
        transcriptions = []
        for start_idx in range(0, len(speech), samples_per_chunk):
            chunk = speech[start_idx : start_idx + samples_per_chunk]
            input_values = tokenizer(chunk, return_tensors='pt').input_values
            logits = model(input_values).logits

            # Decode predicted token IDs
            predicted_ids = torch.argmax(logits, dim=-1)
            transcription = tokenizer.decode(predicted_ids[0])
            transcriptions.append(transcription)

        full_transcription = " ".join(transcriptions)
        return full_transcription
    
    except Exception as e:
        return f"Error during transcription: {str(e)}"
    
def summarize_video(video_path: str) -> str:
    """
    Enhanced video summarization with improved frame analysis, better caption grouping,
    and more coherent narrative generation for higher quality summaries.
    
    :param video_path: Path to the video file.
    :return: A descriptive summary of the video content or an error message.
    """
    # Suppress backend logs unless in debug mode
    debug_mode = False
    if not debug_mode:
        logging.getLogger('transformers').setLevel(logging.ERROR)
        logging.getLogger('PIL').setLevel(logging.ERROR)
        logging.getLogger('matplotlib').setLevel(logging.ERROR)
    
    # Initialize all variables and flags at the beginning
    has_valid_audio = False
    has_transcription = False
    audio_transcription = ""
    summary = ""
    frame_descriptions = []
    
    try:
        if not os.path.exists(video_path):
            return "Error: Video file not found. Please check the file path."
            
        # Load pre-trained vision-language model for image captioning
        model = VisionEncoderDecoderModel.from_pretrained("nlpconnect/vit-gpt2-image-captioning")
        feature_extractor = ViTImageProcessor.from_pretrained("nlpconnect/vit-gpt2-image-captioning")
        tokenizer = AutoTokenizer.from_pretrained("nlpconnect/vit-gpt2-image-captioning")
        
        # Set device to GPU if available
        device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
        model.to(device)
        
        # Open the video file
        cap = cv2.VideoCapture(video_path)
        if not cap.isOpened():
            return "Error: Could not open video file."
            
        # Get video properties
        fps = cap.get(cv2.CAP_PROP_FPS)
        frame_count = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
        duration = frame_count / fps
        
        # Enhanced frame sampling for more detail - sample every second for short videos
        if duration > 300:  # > 5 minutes
            sample_interval = 5  # Every 5 seconds (increased frequency)
        elif duration > 60:  # > 1 minute
            sample_interval = 2  # Every 2 seconds (increased frequency)
        else:
            sample_interval = 1  # Every second for short videos
            
        # Calculate frames to sample with increased minimum and maximum
        frames_to_sample = int(duration / sample_interval)
        frames_to_sample = max(20, min(frames_to_sample, 100))  # Between 20 and 100 frames (increased from 10-50)
        
        # Calculate frame interval
        frame_interval = max(1, int(frame_count / frames_to_sample))
        
        print(f"Video duration: {int(duration)} seconds, sampling {frames_to_sample} frames at interval of {frame_interval} frames")
        
        # Extract frames and generate captions
        frame_images = []  # Store frame images for scene change detection
        current_frame = 0
        
        print(f"Analyzing video... (extracting {frames_to_sample} frames)")
        
        # Scene change detection threshold
        scene_change_threshold = 0.5
        prev_hist = None
        
        # Process frames with improved density and detail capture
        while current_frame < frame_count:
            cap.set(cv2.CAP_PROP_POS_FRAMES, current_frame)
            ret, frame = cap.read()
            
            if not ret:
                break
                
            # Store frame for scene change detection
            frame_images.append(frame)
            
            # Calculate histogram for scene change detection
            hist = cv2.calcHist([frame], [0, 1, 2], None, [8, 8, 8], [0, 256, 0, 256, 0, 256])
            hist = cv2.normalize(hist, hist).flatten()
            
            # Check if this is a scene change or the first frame
            is_scene_change = False
            if prev_hist is None:
                is_scene_change = True  # First frame
            else:
                # Compare histograms to detect scene changes
                diff = cv2.compareHist(prev_hist, hist, cv2.HISTCMP_CORREL)
                if diff < scene_change_threshold:
                    is_scene_change = True
            
            prev_hist = hist
            
            # Process every frame according to our sampling rate (more frequent now)
            # Always process scene changes regardless of interval
            if is_scene_change or current_frame % frame_interval == 0:
                # Convert BGR to RGB (OpenCV uses BGR by default)
                rgb_frame = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
                
                # Prepare image for the model
                pixel_values = feature_extractor(images=rgb_frame, return_tensors="pt").pixel_values.to(device)
                
                # Generate caption with more diverse beam search
                with torch.no_grad():
                    output_ids = model.generate(
                        pixel_values, 
                        max_length=50, 
                        num_beams=5,
                        num_beam_groups=5,
                        diversity_penalty=0.5,  # Encourage diverse captions
                        no_repeat_ngram_size=2  # Avoid repetition
                    )
                    caption = tokenizer.decode(output_ids[0], skip_special_tokens=True)
                
                # Improve caption quality by ensuring it starts with a capital letter and ends with punctuation
                caption = caption[0].upper() + caption[1:]
                if not caption.endswith(('.', '!', '?')):
                    caption += '.'
                    
                # Filter low-confidence visual captions with more lenient thresholds
                # Check for blurry images with adjusted threshold
                laplacian_var = cv2.Laplacian(frame, cv2.CV_64F).var()
                is_blurry = laplacian_var < 80  # More lenient threshold for blurriness
                
                # Check for generic/irrelevant captions with expanded list
                generic_phrases = ["a picture of", "there is", "this is a", "this image shows", "a photo of"]
                is_generic = any(phrase in caption.lower() for phrase in generic_phrases) and len(caption.split()) < 5
                
                # Skip this frame if it's too blurry or the caption is too generic
                if is_blurry or is_generic:
                    if debug_mode:
                        print(f"Skipping low-quality frame: Blurry={is_blurry}, Generic={is_generic}")
                    # Increment frame counter and continue
                    current_frame += max(1, int(frame_interval / 2))  # Skip ahead but less than normal interval
                    continue
                
                # Add timestamp information
                timestamp = current_frame / fps
                minutes, seconds = divmod(int(timestamp), 60)
                timestamp_str = f"{minutes:02d}:{seconds:02d}"
                
                # Mark scene changes in the description
                if is_scene_change and len(frame_descriptions) > 0:
                    scene_marker = "[SCENE CHANGE] "
                else:
                    scene_marker = ""
                
                # Add the caption to our frame descriptions
                frame_descriptions.append(f"[{timestamp_str}] {scene_marker}{caption}")
                
                # Debug output to track frame processing
                if debug_mode:
                    print(f"Processed frame at {timestamp_str}: {caption}")
            
            # Increment the frame counter to process the next frame
            # Use smaller increments for more detailed analysis
            current_frame += frame_interval
        
        # Release video capture
        cap.release()
        
        # Extract audio from video and transcribe it
        print("Extracting and transcribing audio from video...")
        
        # Create a temporary audio file
        temp_audio_path = os.path.splitext(video_path)[0] + "_temp_audio.wav"
        
        # Try multiple methods to extract and transcribe audio
        try:
            # First try with moviepy
            try:
                from moviepy.editor import VideoFileClip
                
                # Extract audio using moviepy
                video_clip = VideoFileClip(video_path)
                
                # Check if video has audio track
                if video_clip.audio is None:
                    audio_transcription = "The video does not contain an audio track."
                    has_valid_audio = False
                else:
                    # Check audio levels to detect muted tracks
                    audio_clip = video_clip.audio
                    try:
                        # Sample audio to check levels
                        audio_array = audio_clip.to_soundarray()
                        audio_level = np.abs(audio_array).mean()
                        
                        if audio_level < 0.01:  # Very low audio level threshold
                            audio_transcription = "The video contains an audio track, but the volume is too low to transcribe."
                            has_valid_audio = False
                        else:
                            # Write audio to file with no logging output
                            audio_clip.write_audiofile(temp_audio_path, logger=None, verbose=False)
                            has_valid_audio = os.path.exists(temp_audio_path) and os.path.getsize(temp_audio_path) > 0
                    except Exception as audio_error:
                        logging.error(f"Audio processing error with moviepy: {audio_error}")
                        has_valid_audio = False
                        
                # Close clips
                if hasattr(video_clip, 'audio') and video_clip.audio is not None:
                    if hasattr(video_clip.audio, 'close'):
                        video_clip.audio.close()
                video_clip.close()
                
            except Exception as moviepy_error:
                logging.error(f"Moviepy audio extraction error: {moviepy_error}")
                has_valid_audio = False
                
                # Try alternative method with ffmpeg if moviepy fails
                try:
                    import subprocess
                    subprocess.run(
                        ["ffmpeg", "-i", video_path, "-q:a", "0", "-map", "a", temp_audio_path],
                        check=True, 
                        stdout=subprocess.PIPE, 
                        stderr=subprocess.PIPE
                    )
                    has_valid_audio = os.path.exists(temp_audio_path) and os.path.getsize(temp_audio_path) > 0
                except Exception as ffmpeg_error:
                    logging.error(f"FFmpeg audio extraction error: {ffmpeg_error}")
                    has_valid_audio = False
                    
                    # Try with pydub as last resort
                    try:
                        video = AudioSegment.from_file(video_path)
                        video.export(temp_audio_path, format="wav")
                        has_valid_audio = os.path.exists(temp_audio_path) and os.path.getsize(temp_audio_path) > 0
                    except Exception as pydub_error:
                        logging.error(f"Pydub audio extraction error: {pydub_error}")
                        has_valid_audio = False
            
            # Transcribe audio if we have a valid audio file
            if has_valid_audio:
                # First try with Wav2Vec2
                try:
                    audio_transcription = transcribe_audio(temp_audio_path)
                    has_transcription = audio_transcription and not audio_transcription.startswith("Error") and len(audio_transcription) > 10
                except Exception as transcribe_error:
                    logging.error(f"Wav2Vec2 transcription error: {transcribe_error}")
                    has_transcription = False
                
                # If Wav2Vec2 fails, try Whisper
                if not has_transcription:
                    try:
                        whisper_model = whisper.load_model("base")
                        result = whisper_model.transcribe(temp_audio_path)
                        audio_transcription = result["text"]
                        has_transcription = audio_transcription and len(audio_transcription) > 10
                    except Exception as whisper_error:
                        logging.error(f"Whisper transcription error: {whisper_error}")
                        has_transcription = False
                        audio_transcription = "Could not transcribe audio from the video."
            else:
                audio_transcription = "No valid audio could be extracted from the video."
                has_transcription = False
                
            # Clean up the temporary audio file
            if os.path.exists(temp_audio_path):
                os.remove(temp_audio_path)
                
        except Exception as e:
            logging.error(f"Audio extraction error: {e}")
            audio_transcription = "Could not extract or transcribe audio from the video."
            has_valid_audio = False
            has_transcription = False
            
            # Clean up the temporary audio file if it exists
            if os.path.exists(temp_audio_path):
                os.remove(temp_audio_path)
        
        # Combine frame descriptions into a coherent summary
        if not frame_descriptions:
            return "Could not extract any meaningful content from the video."
            
        # Generate a summary from the frame descriptions
        descriptions_text = "\n".join(frame_descriptions)
        
        # Use the existing summarization function to create a concise summary
        print("\nFrame-by-frame descriptions:")
        print(descriptions_text)
        
        print("\nGenerating final video summary...")
        
        # Create a summary that captures the essence of the video
        summary = f"Video Summary ({int(duration)} seconds):\n\n"
        
        # Extract captions without timestamps and scene change markers
        captions = []
        for desc in frame_descriptions:
            # Split by timestamp bracket
            parts = desc.split('] ', 1)
            if len(parts) > 1:
                caption = parts[1]
                # Remove scene change marker if present
                caption = caption.replace("[SCENE CHANGE] ", "")
                captions.append(caption)
        
        # Enhanced caption grouping using text similarity
        from sklearn.feature_extraction.text import TfidfVectorizer
        from sklearn.metrics.pairwise import cosine_similarity
        import numpy as np
        
        # Create TF-IDF vectors for all captions
        vectorizer = TfidfVectorizer(stop_words='english')
        try:
            # Only proceed with TF-IDF if we have enough captions
            if len(captions) > 1:
                tfidf_matrix = vectorizer.fit_transform(captions)
                
                # Calculate similarity between all pairs of captions
                similarity_matrix = cosine_similarity(tfidf_matrix)
                
                # Group captions based on similarity threshold
                similarity_threshold = 0.3
                grouped_indices = []
                unprocessed = set(range(len(captions)))
                
                while unprocessed:
                    # Get the first unprocessed caption
                    current_idx = min(unprocessed)
                    current_group = [current_idx]
                    unprocessed.remove(current_idx)
                    
                    # Find all similar captions
                    for idx in list(unprocessed):
                        if similarity_matrix[current_idx, idx] > similarity_threshold:
                            current_group.append(idx)
                            unprocessed.remove(idx)
                    
                    grouped_indices.append(sorted(current_group))
                
                # Convert indices to caption groups
                grouped_captions = [[captions[idx] for idx in group] for group in grouped_indices]
            else:
                # If only one caption, just use it directly
                grouped_captions = [[captions[0]]] if captions else []
            
        except Exception as e:
            logging.error(f"Error in caption grouping: {e}")
            # Fallback to simpler grouping method if TF-IDF fails
            grouped_captions = []
            if captions:
                current_group = [captions[0]]
                
                for i in range(1, len(captions)):
                    # Simple similarity check - if captions share significant words
                    current_words = set(current_group[-1].lower().split())
                    next_words = set(captions[i].lower().split())
                    common_words = current_words.intersection(next_words)
                    
                    # If there's significant overlap, group them
                    if len(common_words) >= 2 and len(common_words) / len(next_words) > 0.3:
                        current_group.append(captions[i])
                    else:
                        # Process the current group before starting a new one
                        if current_group:
                            grouped_captions.append(current_group)
                        current_group = [captions[i]]
                
                # Add the last group
                if current_group:
                    grouped_captions.append(current_group)
        
        # Generate coherent paragraphs from grouped captions with improved narrative flow
        paragraphs = []
        
        # Enhanced transitional phrases organized by narrative position
        intro_transitions = [
            "The video begins with", "Initially, the video shows", "At the start,", 
            "The video opens with", "The footage starts by showing"
        ]
        
        middle_transitions = [
            "The scene then changes to", "Subsequently,", "Following this,", 
            "Next, we can see", "The video continues with", "Later,", 
            "Towards the middle,", "As the video progresses,", 
            "The focus then shifts to", "Afterward,", "Then,",
            "This is followed by", "The next segment shows", "Moving forward,"
        ]
        
        conclusion_transitions = [
            "Towards the end,", "Finally,", "The video concludes with",
            "In the closing scene,", "The final segment displays", "To conclude,"
        ]
        
        # Determine narrative position for each group
        total_groups = len(grouped_captions)
        
        for i, group in enumerate(grouped_captions):
            # Choose an appropriate transition based on narrative position
            if i == 0:
                # Introduction - randomly select from intro transitions
                import random
                transition = random.choice(intro_transitions)
            elif i == total_groups - 1:
                # Conclusion - randomly select from conclusion transitions
                import random
                transition = random.choice(conclusion_transitions)
            else:
                # Middle section - select transition based on relative position
                position_ratio = i / total_groups
                # Use different transitions based on whether we're in early, middle, or late middle
                if position_ratio < 0.33:
                    idx = i % (len(middle_transitions) // 3)
                elif position_ratio < 0.66:
                    idx = (len(middle_transitions) // 3) + (i % (len(middle_transitions) // 3))
                else:
                    idx = (2 * len(middle_transitions) // 3) + (i % (len(middle_transitions) // 3))
                
                transition = middle_transitions[min(idx, len(middle_transitions)-1)]
            
            # Improved caption combination for more natural language
            if len(group) == 1:
                paragraph = f"{transition} {group[0]}"
            else:
                # Use NLP techniques to combine similar captions more naturally
                # Extract key subjects, actions, and objects from the captions
                try:
                    # Use spaCy for better linguistic analysis
                    doc = nlp(" ".join(group))
                    
                    # Extract main subjects and objects
                    subjects = []
                    actions = []
                    objects = []
                    
                    for token in doc:
                        if token.dep_ in ("nsubj", "nsubjpass"):
                            subjects.append(token.text)
                        elif token.pos_ == "VERB":
                            actions.append(token.text)
                        elif token.dep_ in ("dobj", "pobj"):
                            objects.append(token.text)
                    
                    # Create a more natural combined description
                    if subjects and actions:
                        main_subject = subjects[0]
                        main_action = actions[0]
                        main_object = objects[0] if objects else ""
                        
                        combined = f"{main_subject} {main_action}"
                        if main_object:
                            combined += f" {main_object}"
                            
                        # Add additional details from other captions
                        additional_details = []
                        for caption in group[1:]:
                            # Extract unique phrases not in the main description
                            caption_doc = nlp(caption)
                            for chunk in caption_doc.noun_chunks:
                                if chunk.text.lower() not in combined.lower():
                                    additional_details.append(chunk.text)
                        
                        if additional_details:
                            combined += f" with {', '.join(additional_details[:2])}"
                    else:
                        # Fallback to simpler combination
                        combined = group[0]
                        for caption in group[1:]:
                            # Find unique elements in this caption
                            current_words = set(combined.lower().split())
                            new_words = set(caption.lower().split())
                            unique_words = new_words - current_words
                            
                            if unique_words:
                                # Add unique elements to the combined description
                                unique_phrase = ' '.join([w for w in caption.split() if w.lower() in unique_words])
                                combined += f" with {unique_phrase}"
                    
                    paragraph = f"{transition} {combined}"
                except Exception as e:
                    logging.error(f"Error in caption combination: {e}")
                    # Fallback to simple combination
                    combined = group[0]
                    paragraph = f"{transition} {combined}"
            
            paragraphs.append(paragraph)
        
        # Combine paragraphs into a flowing narrative with improved coherence
        narrative = ". ".join(paragraphs) + "."
        
        # Clean up any double periods or spacing issues
        narrative = narrative.replace("..", ".").replace(" .", ".").replace(".", ". ").strip()
        
        # Apply additional post-processing for readability
        # Fix any remaining spacing issues
        narrative = re.sub(r'\s+', ' ', narrative)
        # Ensure proper capitalization after periods
        narrative = re.sub(r'\. ([a-z])', lambda m: f". {m.group(1).upper()}", narrative)
        
        # Post-process with GPT or T5 for grammar correction and improved readability
        try:
            # Use a pre-trained T5 model for grammar correction
            grammar_model = pipeline("text2text-generation", model="google/t5-small")
            corrected_narrative = grammar_model(f"grammar: {narrative}", max_length=len(narrative.split()) + 50)[0]['generated_text']
            
            # Only use the corrected version if it's not significantly shorter than the original
            if len(corrected_narrative.split()) > len(narrative.split()) * 0.8:
                narrative = corrected_narrative
        except Exception as e:
            logging.error(f"Grammar correction error: {e}")
            # Continue with the uncorrected narrative
        
        # Create a structured summary with introduction, middle, and conclusion
        structured_summary = ""
        
        # Introduction section
        intro = ""
        if has_valid_audio and 'processed_transcription' in locals():
            # Extract the first sentence or two from audio transcription for intro
            audio_sentences = sent_tokenize(processed_transcription)
            if audio_sentences:
                intro = f"Introduction: {audio_sentences[0]}"
                if len(audio_sentences) > 1:
                    intro += f" {audio_sentences[1]}"
        
        if not intro and paragraphs:
            # Use the first paragraph as introduction if no audio
            intro = f"Introduction: {paragraphs[0]}"
        
        # Middle section - main content
        middle = f"Content: {narrative}"
        
        # Conclusion section
        conclusion = ""
        if paragraphs:
            conclusion = f"Conclusion: {paragraphs[-1]}"
        
        # Combine the structured sections
        structured_summary = f"{intro}\n\n{middle}\n\n{conclusion}"
        
        # Make sure we have actual content in the narrative
        if structured_summary.strip():
            summary += structured_summary
        else:
            summary += "The video contains visual content that could not be automatically described."
        
        # Enhanced audio extraction and integration with visual content
        # Check if audio transcription was successful and contains meaningful content
        if 'audio_transcription' in locals():
            # Validate audio transcription content
            if (audio_transcription and 
                not audio_transcription.startswith("Error") and 
                not audio_transcription.startswith("Could not") and
                not audio_transcription.startswith("The video does not contain") and
                not audio_transcription.startswith("The video contains an audio track, but") and
                len(audio_transcription.strip()) > 20):  # Ensure we have meaningful content
                
                has_valid_audio = True
                processed_transcription = audio_transcription.strip()
                if debug_mode:
                    print("\nAudio Transcription:")
                    print(processed_transcription)
            else:
                # Try one more time with ffmpeg directly if available
                try:
                    import subprocess
                    import shutil
                    
                    # Check if ffmpeg is available
                    if shutil.which("ffmpeg"):
                        print("Attempting audio extraction with ffmpeg directly...")
                        # Create a temporary audio file
                        temp_audio_path = os.path.splitext(video_path)[0] + "_ffmpeg_audio.wav"
                        
                        # Extract audio using ffmpeg
                        subprocess.run([
                            "ffmpeg", "-y", "-i", video_path, 
                            "-vn", "-acodec", "pcm_s16le", "-ar", "16000", "-ac", "1",
                            temp_audio_path
                        ], check=True, capture_output=True)
                        
                        # Try transcribing with Whisper first (usually more reliable)
                        try:
                            whisper_model = whisper.load_model("base")
                            result = whisper_model.transcribe(temp_audio_path)
                            processed_transcription = result["text"].strip()
                            
                            if processed_transcription and len(processed_transcription) > 20:
                                has_valid_audio = True
                                print("\nAudio Transcription (ffmpeg + whisper):")
                                print(processed_transcription)
                        except Exception as whisper_error:
                            logging.error(f"Whisper transcription error: {whisper_error}")
                            
                            # Fall back to Wav2Vec2 if Whisper fails
                            try:
                                processed_transcription = transcribe_audio(temp_audio_path)
                                if (processed_transcription and 
                                    not processed_transcription.startswith("Error") and
                                    len(processed_transcription) > 20):
                                    has_valid_audio = True
                                    print("\nAudio Transcription (ffmpeg + wav2vec2):")
                                    print(processed_transcription)
                            except Exception as w2v_error:
                                logging.error(f"Wav2Vec2 transcription error: {w2v_error}")
                        
                        # Clean up temporary file
                        if os.path.exists(temp_audio_path):
                            os.remove(temp_audio_path)
                except Exception as ffmpeg_error:
                    logging.error(f"FFmpeg extraction error: {ffmpeg_error}")
        
        # Add audio transcription to the summary if available
        if has_valid_audio:
            # Add the transcription section to the summary
            summary += "\n\nAudio Transcription:\n\n"
            
            # If the transcription is very long, summarize it
            if len(processed_transcription.split()) > 200:
                # Use the existing summarization function to create a concise summary of the transcription
                print("Summarizing audio transcription...")
                try:
                    # Initialize a summarization model
                    model_name = "facebook/bart-large-cnn"
                    summarizer = pipeline("summarization", model=model_name, tokenizer=model_name, framework="pt", device=0 if torch.cuda.is_available() else -1)
                    
                    # Split text into manageable chunks and summarize each chunk
                    text_chunks = chunk_text(processed_transcription)
                    transcription_summaries = []
                    
                    for chunk in text_chunks:
                        if len(chunk.split()) > 10:  # Only summarize if there's enough content
                            summary_output = summarizer(chunk, max_length=100, min_length=30, do_sample=False)
                            transcription_summaries.append(summary_output[0]['summary_text'])
                        else:
                            transcription_summaries.append(chunk)
                    
                    # Combine summaries
                    processed_transcription = " ".join(transcription_summaries)
                except Exception as e:
                    logging.error(f"Transcription summarization error: {e}")
                    # If summarization fails, use the original transcription
                    pass
            
            summary += processed_transcription
            
            # Create a combined summary section that integrates visual and audio information
            summary += "\n\nIntegrated Summary:\n\n"
            
            # Create a more coherent narrative by combining visual and audio information
            try:
                # Initialize a summarization model for the final integration
                model_name = "facebook/bart-large-cnn"
                integrator = pipeline("summarization", model=model_name, tokenizer=model_name, framework="pt", device=0 if torch.cuda.is_available() else -1)
                
                # Combine visual narrative and processed transcription with clear separation
                combined_text = f"Visual content: {narrative} Audio content: {processed_transcription}"
                
                # Generate an integrated summary
                # Ensure min_length is always less than max_length
                combined_text_length = len(combined_text.split())
                max_length = min(150, combined_text_length)
                min_length = min(30, max_length - 5)  # Ensure min_length is less than max_length
                
                integrated_summary = integrator(combined_text, max_length=max_length, min_length=min_length, do_sample=False)
                
                # Post-process with T5 for paraphrasing and grammar cleanup
                try:
                    # Use a pre-trained T5 model for paraphrasing
                    paraphrase_model = pipeline("text2text-generation", model="google/t5-small")
                    paraphrased_summary = paraphrase_model(
                        f"paraphrase: {integrated_summary[0]['summary_text']}", 
                        max_length=len(integrated_summary[0]['summary_text'].split()) + 50,
                        do_sample=True,
                        top_k=50,
                        top_p=0.95
                    )[0]['generated_text']
                    
                    # Only use the paraphrased version if it's not significantly shorter
                    if len(paraphrased_summary.split()) > len(integrated_summary[0]['summary_text'].split()) * 0.8:
                        summary += paraphrased_summary
                    else:
                        summary += integrated_summary[0]['summary_text']
                except Exception as e:
                    logging.error(f"Paraphrasing error: {e}")
                    # Continue with the original summary
                    summary += integrated_summary[0]['summary_text']
                
                # Validate the integrated summary contains both visual and audio elements
                if "visual" not in summary.lower() and "audio" not in summary.lower():
                    # Add explicit markers if the model didn't include them
                    summary += "\n\nThis summary combines visual elements showing " + narrative.lower() + " with audio that discusses " + processed_transcription[:100] + "..."
            except Exception as e:
                logging.error(f"Integration summarization error: {e}")
                # If integration fails, provide a simple combined summary with clear separation
                summary += "This video contains visual elements showing " + narrative.lower() + "\n\nThe audio discusses: " + processed_transcription[:200] + "..."
        else:
            # No valid audio transcription available
            summary += "\n\nNote: No audio content could be extracted or transcribed from this video."
            
            # Add a note about potential reasons
            summary += "\nPossible reasons: The video may not have an audio track, the audio quality may be poor, or there might be an issue with the audio extraction process."
            
            # Still provide the visual summary
            summary += "\n\nVisual Content Summary:\n\n" + narrative
        
        # Format the complete summary with all extracted information using a structured template
        complete_summary = f"Video Summary ({int(duration)} seconds):\n\n"
        
        # Ensure the summary contains meaningful content from both visual and audio sources when available
        if has_valid_audio:
            # Check if the summary already contains the integrated summary section
            if "Integrated Summary:" not in summary:
                # Add a properly formatted integrated summary section
                visual_summary = narrative
                audio_summary = processed_transcription[:200] + "..." if len(processed_transcription) > 200 else processed_transcription
                
                complete_summary += f"Integrated Summary:\n\nThis video shows {visual_summary.lower()}\n\nThe audio content discusses: {audio_summary}\n\n"
            else:
                complete_summary += summary
        else:
            complete_summary += summary
        
        # Add frame descriptions section
        if frame_descriptions:
            complete_summary += "Frame-by-frame descriptions:\n"
            complete_summary += "\n".join(frame_descriptions) + "\n\n"
        
        # Add audio transcription section if available
        if 'audio_transcription' in locals() and audio_transcription and not audio_transcription.startswith("Error"):
            complete_summary += "Audio content:\n" + audio_transcription[:500] + ("..." if len(audio_transcription) > 500 else "") + "\n\n"
        
        # Add the generated summary section with structured template format
        if summary.strip() in ["", f"Video Summary ({int(duration)} seconds):\n\n", "Watch the video to.", "Watch this video to see."] or len(summary.split()) < 10:
            # Generate a fallback summary from the frame descriptions directly using a structured template
            complete_summary += "Introduction: The video contains visual content.\n\n"
            complete_summary += "Content: "
            clean_descriptions = []
            for desc in frame_descriptions:
                parts = desc.split('] ', 1)
                if len(parts) > 1:
                    clean_desc = parts[1].replace("[SCENE CHANGE] ", "")
                    if clean_desc not in clean_descriptions:
                        clean_descriptions.append(clean_desc)
            
            # Combine the first several unique descriptions
            if clean_descriptions:
                summary = "The video contains: " + ". ".join(clean_descriptions[:min(10, len(clean_descriptions))])
            else:
                summary = "The video shows " + (frame_descriptions[0].split('] ')[1] if frame_descriptions and '] ' in frame_descriptions[0] else "visual content that requires more detailed analysis.")
        
        # Only add the Integrated Summary section if it's not already included
        if "Integrated Summary:" not in complete_summary:
            complete_summary += "Integrated Summary:\n" + summary
        
        # Validate the final summary before returning
        if not complete_summary or len(complete_summary.strip()) < 50:
            # If the summary is too short or empty, provide a fallback summary
            return f"Video Summary ({int(duration)} seconds):\n\nThe video contains visual content that could not be fully analyzed. Duration: {int(duration)} seconds.\n\nNote: The automatic summarization process encountered difficulties extracting meaningful content from this video."
        
        # Final cleanup and formatting
        complete_summary = re.sub(r'\s+', ' ', complete_summary)  # Remove excessive whitespace
        complete_summary = re.sub(r'\n\s*\n', '\n\n', complete_summary)  # Normalize paragraph breaks
        
        if debug_mode:
            print("\nVideo summarization complete!")
            
        return complete_summary
        
    except Exception as e:
        logging.error(f"Video summarization error: {e}")
        return f"Error during video summarization: {str(e)}"

def answer_question(text, ask):
    """
    Answer questions based on the given text using a state-of-the-art QA model with enhanced response quality.
    
    :param text: The original document content.
    :param ask: The question asked by the user.
    :return: Detailed answer extracted from the text.
    """
    try:
        # Always use CPU to avoid CUDA errors
        device = -1  # Force CPU usage
        
        # Initialize the QA pipeline with the latest model
        qa_model = pipeline(
            "question-answering", 
            model="deepset/roberta-base-squad2",
            device=device
        )
        
        # For very long texts, use chunking to process in parts
        if len(text.split()) > 500:
            # Split text into chunks with overlap
            chunks = chunk_text(text, chunk_size=500, overlap=150, respect_sentences=True)
            
            # Process each chunk and collect answers
            answers = []
            scores = []
            contexts = []
            
            for chunk in chunks:
                # Get answer from this chunk
                result = qa_model(
                    question=ask, 
                    context=chunk
                )
                
                # Store answer, score and context
                answers.append(result["answer"])
                scores.append(result["score"])
                
                # Store the sentence containing this answer for context
                for sentence in sent_tokenize(chunk):
                    if result["answer"] in sentence:
                        contexts.append(sentence)
                        break
                else:
                    contexts.append("")  # No context found
            
            # Find the best answer based on confidence score
            if answers:
                best_idx = scores.index(max(scores))
                best_answer = answers[best_idx]
                confidence = scores[best_idx]
                context = contexts[best_idx]
                
                # Provide more context if confidence is high but answer is short
                if confidence > 0.7 and len(best_answer.split()) < 5 and context:
                    return f"{best_answer} (Context: {context})"
                elif confidence > 0.9:
                    return best_answer
                elif confidence > 0.5:
                    return best_answer
                else:
                    return f"{best_answer} (Note: Low confidence answer)"
            else:
                return "No answer found in the document."
        else:
            # For shorter texts, process directly
            result = qa_model(
                question=ask, 
                context=text
            )
            
            # Process and return the best answer
            if result["score"] > 0.7:
                return result["answer"]
            elif result["score"] > 0.5:
                # Find the sentence containing this answer for context
                for sentence in sent_tokenize(text):
                    if result["answer"] in sentence:
                        return f"{result['answer']} (Context: {sentence})"
                return result["answer"]
            else:
                return f"{result['answer']} (Note: Low confidence answer)"
                
    except Exception as e:
        logging.error(f"Question answering error: {e}")
        # Try a simpler fallback approach with DistilBERT model
        try:
            # Fallback to a simpler model if available
            fallback_model = pipeline(
                "question-answering", 
                model="distilbert-base-cased-distilled-squad",
                device=-1  # Force CPU for fallback
            )
            
            result = fallback_model(
                question=ask, 
                context=text[:5000]  # Use first 5000 chars for fallback
            )
            
            return f"{result['answer']} (Fallback answer)"
        except Exception as fallback_error:
            logging.error(f"Fallback question answering error: {fallback_error}")
            return f"Error processing question. Please try again with a simpler question or shorter text."


# Output Functions
def save_as_txt(text, filename="summary_output.txt"):
    # Check if the text is from video summarization (contains frame descriptions)
    if "[SCENE CHANGE]" in text or "[00:" in text:
        # Extract frame descriptions if they exist in the text
        frame_descriptions = [line for line in text.split('\n') if line.startswith('[') and ':' in line.split(']')[0]]
        
        # Extract any audio transcription if present
        audio_content = ""
        if "Audio content:" in text:
            audio_parts = text.split("Audio content:", 1)
            if len(audio_parts) > 1:
                audio_content = "Audio content: " + audio_parts[1].strip()
        
        # Extract copyright notice if present
        copyright_notice = ""
        if "All images subject to copyright" in text:
            # Find the copyright notice and extract it
            for line in text.split('\n'):
                if "All images subject to copyright" in line:
                    copyright_notice = line.strip()
                    break
        
        # Check if we have an actual summary beyond just copyright and a single frame description
        has_meaningful_content = len(frame_descriptions) > 1 or audio_content or "Integrated Summary:" in text
        
        # If we don't have meaningful content, try to generate a basic summary from what we have
        if not has_meaningful_content and frame_descriptions:
            # Extract the actual description part from the frame description
            description_parts = frame_descriptions[0].split('] ', 1)
            if len(description_parts) > 1:
                description_text = description_parts[1].replace("[SCENE CHANGE] ", "")
                # Create a basic summary from this description
                basic_summary = f"The video shows {description_text}"
                # Add this as an integrated summary
                text += "\n\nIntegrated Summary:\n" + basic_summary
                # Also add it to our tracking variable so we know we have a summary now
                has_meaningful_content = True
        
        # Format the complete summary with all extracted information
        complete_summary = "Video Summary:\n\n"
        
        # Add copyright notice if available
        if copyright_notice:
            complete_summary += copyright_notice + "\n\n"
        
        # Add frame descriptions
        if frame_descriptions:
            complete_summary += "Frame-by-frame descriptions:\n"
            complete_summary += "\n".join(frame_descriptions) + "\n\n"
        
        # Add audio transcription if available
        if audio_content:
            complete_summary += audio_content + "\n\n"
        
        # Add the generated summary
        if "Integrated Summary:" in text:
            summary_parts = text.split("Integrated Summary:", 1)
            if len(summary_parts) > 1:
                complete_summary += "Integrated Summary:\n" + summary_parts[1].strip()
        else:
            # If no integrated summary, use the original text but exclude the copyright notice
            # to avoid duplication since we already added it at the beginning
            filtered_text = text
            if copyright_notice and copyright_notice in filtered_text:
                filtered_text = '\n'.join([line for line in filtered_text.split('\n') 
                                          if line.strip() != copyright_notice])
            complete_summary += "Summary:\n" + filtered_text
        
        # Write the complete summary to file
        with open(filename, "w", encoding="utf-8") as f:
            f.write(complete_summary)
    else:
        # For non-video summaries, just write the text as is
        with open(filename, "w", encoding="utf-8") as f:
            f.write(text)

def save_as_json(data, filename="summary_output.json"):
    with open(filename, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=4)

def save_as_pdf(text, filename="summary_output.pdf"):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, text)
    pdf.output(filename)

def save_as_docx(text, filename="summary_output.docx"):
    doc = Document()
    doc.add_paragraph(text)
    doc.save(filename)

# Main Application
def main():
    print("Welcome to the Advanced Text Summarizer!")
    print("Choose a summarization model (t5-small/bart-large/distilbart/pegasus/t5-3b):")
    model_choice = input("Model: ")
    user_input = ""

    while True:
        print("\nOptions: text, file, pdf, image, audio, video, exit")
        choice = input("Enter your choice: ").lower()

        if choice == 'exit':
            print("Exiting. Goodbye!")
            break

        elif choice == 'text':
            user_input = input("Enter text: ")
        
        elif choice == 'file':
            filename = input("Enter filename (.txt): ")
            with open(filename, "r", encoding="utf-8") as f:
                user_input = f.read()

        elif choice == 'pdf':
            filename = input("Enter PDF filename: ")
            user_input = read_pdf(filename)

        elif choice == 'image':
            filename = input("Enter image filename: ")
            user_input = extract_text_from_image(filename)
            if "Error" in user_input:
                print(user_input)
                continue


        elif choice == 'audio':
            filename = input("Enter audio filename: ")
            user_input = transcribe_audio(filename)
            
        elif choice == 'video':
            filename = input("Enter video filename: ")
            user_input = summarize_video(filename)
            if "Error" in user_input:
                print(user_input)
                continue
        
        else:
            print("Invalid choice.")
            continue

        if not user_input.strip():
            print("Empty input! Try again.")
            continue

        language = detect(user_input)
        if language != "en":
            print(f"Detected language: {language}, translating to English...")
            user_input = GoogleTranslator(source=language, target="en").translate(user_input)


        processed_text = preprocess_text(user_input)
        summary = summarize_text(processed_text, model_choice)


        print("\nSummary:")
        print(summary)

        save_format = input("Save as (txt/json/pdf/docx): ").lower()
        if save_format == 'txt':
            save_as_txt(summary)
        elif save_format == 'json':
            save_as_json({"summary": summary})
        elif save_format == 'pdf':
            save_as_pdf(summary)
        elif save_format == 'docx':
            save_as_docx(summary)
        else:
            print("Invalid save format.")
        
        while True:
            ask_choice = input("Do you want to ask a question about the text? (yes/no): ").lower()
            if ask_choice == 'yes':
                ask = input("Enter your question: ")
                answer = answer_question(user_input, ask)
                print("\nAnswer:\n", answer)
            elif ask_choice == 'no':
                break
            else:
                print("Invalid input! Please type 'yes' or 'no'.")

if __name__ == "__main__":
    main()
=======
import re
import json
import logging
import nltk
from transformers import pipeline
from textwrap import wrap
from langdetect import detect
from nltk.corpus import stopwords
from textblob import TextBlob
from deep_translator import GoogleTranslator
from nltk.tokenize import sent_tokenize, word_tokenize
from rake_nltk import Rake
import spacy
import PyPDF2
import pytesseract
from PIL import Image
import whisper
from fpdf import FPDF
from docx import Document
import textstat
import easyocr 
import speech_recognition as sr
import torch
import openai
import os
from pydub import AudioSegment
import librosa
import torch
from transformers import Wav2Vec2ForCTC, Wav2Vec2Tokenizer
import cv2
from transformers import VisionEncoderDecoderModel, ViTImageProcessor, AutoTokenizer


# Initialize necessary components
logging.basicConfig(filename="summarizer.log", level=logging.ERROR, format="%(asctime)s - %(levelname)s - %(message)s")
models = {
    "t5-small": "t5-small",
    "bart-large": "facebook/bart-large-cnn",
    "distilbart": "sshleifer/distilbart-cnn-12-6",
    "pegasus": "google/pegasus-xsum",
    "t5-3b": "t5-3b"
}
stopwords.words("english")  # Download if not already done
nlp = spacy.load("en_core_web_sm")

# Preprocessing Functions
def preprocess_text(text):
    text = text.lower().strip()
    text = re.sub(r'[^a-zA-Z0-9\s]', '', text)  # Allow only alphanumeric and spaces
    text = re.sub(r'\s+', ' ', text).strip()  # Remove extra spaces
    words = word_tokenize(text)
    words = [word for word in words if word.lower() not in stopwords.words("english") and len(word) > 2]

    return " ".join(words)

def chunk_text(text, chunk_size=500):
    return wrap(text, width=chunk_size)

# Summarization Functions
def summarize_text(text, model_choice):
    try:
        # Get the maximum word limit from user input
        max_words = int(input("Enter maximum word limit for the summary (recommended: 20-30% of original length): "))
        
        # Initialize the summarization model
        model_name = models.get(model_choice, "facebook/bart-large-cnn")
        model = pipeline("summarization", model=model_name, tokenizer=model_name, framework="pt", device=0 if torch.cuda.is_available() else -1)

        
        # Split text into manageable chunks and summarize each chunk
        text_chunks = chunk_text(text)
        summaries = []
        for chunk in text_chunks:
            try:
                summary_output = model(chunk, max_length=max_words, min_length=10, do_sample=False)
                summaries.append(summary_output[0]['summary_text'])
            except Exception as e:
                logging.error(f"Summarization error: {e}")
                summaries.append(chunk)  # Fallback to original chunk if summarization fails


        
        # Combine summaries and truncate to match the max word limit
        final_summary = " ".join(summaries)
        final_summary_words = final_summary.split()
        if len(final_summary_words) > max_words:
            final_summary = " ".join(final_summary_words[:max_words])
        
        return final_summary
    except Exception as e:
        logging.error(f"Summarization error: {e}")
        return "Error in summarization. Check logs."


# Additional Features
def extract_keywords(text):
    r = Rake()
    r.extract_keywords_from_text(text)
    return r.get_ranked_phrases()

def extract_entities(text):
    doc = nlp(text)
    return [(ent.text, ent.label_) for ent in doc.ents]

def analyze_sentiment(text):
    return TextBlob(text).sentiment.polarity

def detect_bias(text):
    classifier = pipeline("text-classification", model="facebook/bart-large-mnli")
    labels = ["biased", "neutral"]
    result = classifier(text, candidate_labels=labels)
    return result

def classify_topic(text):
    classifier = pipeline("zero-shot-classification")
    labels = ["Finance", "Health", "Technology", "Education"]
    result = classifier(text, candidate_labels=labels)
    return result["labels"][0]

def compression_ratio(original, summary):
    return (1 - len(summary) / len(original)) * 100

def readability_score(text):
    return round(textstat.flesch_reading_ease(text), 2)  # Round for consistency


# File Handling Functions
def read_pdf(file_path):
    try:
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            return "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
    except Exception as e:
        logging.error(f"PDF reading error: {e}")
        return "Error reading PDF file."

import easyocr
import logging
import os

# Configure logging
logging.basicConfig(level=logging.WARNING, format="%(levelname)s: %(message)s")

def extract_text_from_image(image_path):
    try:
        # Initialize EasyOCR reader (English language, using GPU if available)
        reader = easyocr.Reader(['en'], gpu=True)

        # Perform OCR
        results = reader.readtext(image_path, detail=1)  # Returns text along with confidence

        
        

        extracted_text = []

        for entry in results:
            if len(entry) >= 2:  # Ensure structure contains text and confidence
                text, confidence = entry[1], entry[2]
                
                # Apply confidence threshold (e.g., 60%)
                if confidence > 0.6:
                    extracted_text.append(text)

        final_text = "\n".join(extracted_text)  # Join text with new lines

        return final_text if final_text.strip() else "No text detected."

    except Exception as e:
        return f"Error extracting text from image: {str(e)}"

def transcribe_audio(file_path: str) -> str:
    """
    Transcribe an audio file using Wav2Vec2 model for speech recognition.

    :param file_path: Path to the audio file (WAV or other formats like MP3).
    :return: Transcribed text as a string or an error message.
    """
    try:
        if not os.path.exists(file_path):
            return "Error: Audio file not found. Please check the file path."

        # Load pre-trained model and tokenizer
        tokenizer = Wav2Vec2Tokenizer.from_pretrained("facebook/wav2vec2-base-960h")
        model = Wav2Vec2ForCTC.from_pretrained("facebook/wav2vec2-base-960h")
        
        # Load audio
        sample_rate = 16000
        speech, rate = librosa.load(file_path, sr=sample_rate)
        
        # Define chunk size (30 seconds)
        chunk_duration = 30
        samples_per_chunk = chunk_duration * sample_rate

        # Process and transcribe each chunk
        transcriptions = []
        for start_idx in range(0, len(speech), samples_per_chunk):
            chunk = speech[start_idx : start_idx + samples_per_chunk]
            input_values = tokenizer(chunk, return_tensors='pt').input_values
            logits = model(input_values).logits

            # Decode predicted token IDs
            predicted_ids = torch.argmax(logits, dim=-1)
            transcription = tokenizer.decode(predicted_ids[0])
            transcriptions.append(transcription)

        full_transcription = " ".join(transcriptions)
        return full_transcription
    
    except Exception as e:
        return f"Error during transcription: {str(e)}"
    
def summarize_video(video_path: str) -> str:
    """
    Analyze and summarize video content by extracting frames, generating descriptions,
    and transcribing audio to create a comprehensive summary.
    
    :param video_path: Path to the video file.
    :return: A descriptive summary of the video content or an error message.
    """
    try:
        if not os.path.exists(video_path):
            return "Error: Video file not found. Please check the file path."
            
        # Load pre-trained vision-language model for image captioning
        model = VisionEncoderDecoderModel.from_pretrained("nlpconnect/vit-gpt2-image-captioning")
        feature_extractor = ViTImageProcessor.from_pretrained("nlpconnect/vit-gpt2-image-captioning")
        tokenizer = AutoTokenizer.from_pretrained("nlpconnect/vit-gpt2-image-captioning")
        
        # Set device to GPU if available
        device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
        model.to(device)
        
        # Open the video file
        cap = cv2.VideoCapture(video_path)
        if not cap.isOpened():
            return "Error: Could not open video file."
            
        # Get video properties
        fps = cap.get(cv2.CAP_PROP_FPS)
        frame_count = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
        duration = frame_count / fps
        
        # Calculate frame sampling rate (extract frames every N seconds)
        sample_interval = 5  # Extract a frame every 5 seconds
        frames_to_sample = int(duration / sample_interval)
        frames_to_sample = min(frames_to_sample, 20)  # Limit to 20 frames maximum
        
        if frames_to_sample <= 0:
            frames_to_sample = 1  # Ensure at least one frame is processed
            
        frame_interval = max(1, int(frame_count / frames_to_sample))
        
        # Extract frames and generate captions
        frame_descriptions = []
        current_frame = 0
        
        print(f"Analyzing video... (extracting {frames_to_sample} frames)")
        
        while current_frame < frame_count:
            cap.set(cv2.CAP_PROP_POS_FRAMES, current_frame)
            ret, frame = cap.read()
            
            if not ret:
                break
                
            # Convert BGR to RGB (OpenCV uses BGR by default)
            rgb_frame = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
            
            # Prepare image for the model
            pixel_values = feature_extractor(images=rgb_frame, return_tensors="pt").pixel_values.to(device)
            
            # Generate caption
            with torch.no_grad():
                output_ids = model.generate(pixel_values, max_length=50, num_beams=4)
                caption = tokenizer.decode(output_ids[0], skip_special_tokens=True)
            
            # Add timestamp information
            timestamp = current_frame / fps
            minutes, seconds = divmod(int(timestamp), 60)
            timestamp_str = f"{minutes:02d}:{seconds:02d}"
            
            frame_descriptions.append(f"[{timestamp_str}] {caption}")
            current_frame += frame_interval
        
        # Release video capture
        cap.release()
        
        # Extract audio from video and transcribe it
        print("Extracting and transcribing audio from video...")
        
        # Create a temporary audio file
        temp_audio_path = os.path.splitext(video_path)[0] + "_temp_audio.wav"
        
        try:
            # Use moviepy to extract audio from video
            from moviepy.editor import VideoFileClip # type: ignore
            
            # Extract audio using moviepy
            print("Extracting audio using moviepy...")
            video_clip = VideoFileClip(video_path)
            audio_clip = video_clip.audio
            audio_clip.write_audiofile(temp_audio_path, logger=None)
            audio_clip.close()
            video_clip.close()
            
            # Transcribe the extracted audio using Wav2Vec2
            audio_transcription = transcribe_audio(temp_audio_path)
            
            # If Wav2Vec2 transcription fails or returns an error, try using Whisper as a fallback
            if audio_transcription.startswith("Error") or not audio_transcription.strip():
                print("Trying alternative transcription method with Whisper...")
                try:
                    # Load Whisper model
                    whisper_model = whisper.load_model("base")
                    
                    # Transcribe with Whisper
                    result = whisper_model.transcribe(temp_audio_path)
                    audio_transcription = result["text"]
                except Exception as whisper_error:
                    logging.error(f"Whisper transcription error: {whisper_error}")
                    # Keep the original error message if Whisper also fails
            
            # Clean up the temporary audio file
            if os.path.exists(temp_audio_path):
                os.remove(temp_audio_path)
                
        except Exception as e:
            logging.error(f"Audio extraction error: {e}")
            audio_transcription = "Could not extract or transcribe audio from the video."
            
            # Clean up the temporary audio file if it exists
            if os.path.exists(temp_audio_path):
                os.remove(temp_audio_path)
        
        # Combine frame descriptions into a coherent summary
        if not frame_descriptions:
            return "Could not extract any meaningful content from the video."
            
        # Generate a summary from the frame descriptions
        descriptions_text = "\n".join(frame_descriptions)
        
        # Use the existing summarization function to create a concise summary
        print("\nFrame-by-frame descriptions:")
        print(descriptions_text)
        
        print("\nGenerating final video summary...")
        
        # Create a summary that captures the essence of the video
        summary = f"Video Summary ({int(duration)} seconds):\n\n"
        
        # Extract captions without timestamps
        captions = [desc.split('] ')[1] for desc in frame_descriptions]
        
        # Group similar captions to avoid repetition
        grouped_captions = []
        current_group = [captions[0]]
        
        for i in range(1, len(captions)):
            # Simple similarity check - if captions share significant words
            current_words = set(current_group[-1].lower().split())
            next_words = set(captions[i].lower().split())
            common_words = current_words.intersection(next_words)
            
            # If there's significant overlap, group them
            if len(common_words) >= 2 and len(common_words) / len(next_words) > 0.3:
                current_group.append(captions[i])
            else:
                # Process the current group before starting a new one
                if current_group:
                    grouped_captions.append(current_group)
                current_group = [captions[i]]
        
        # Add the last group
        if current_group:
            grouped_captions.append(current_group)
        
        # Generate coherent paragraphs from grouped captions
        paragraphs = []
        
        # Transitional phrases to make the summary more natural
        transitions = [
            "The video begins with", "Initially, the video shows", "At the start,", 
            "The scene then changes to", "Subsequently,", "Following this,", 
            "Next, we can see", "The video continues with", "Later,", 
            "Towards the middle,", "As the video progresses,", 
            "The focus then shifts to", "Afterward,", 
            "Towards the end,", "Finally,", "The video concludes with"
        ]
        
        for i, group in enumerate(grouped_captions):
            # Choose an appropriate transition based on position in the video
            if i == 0:
                transition = transitions[0]
            elif i == len(grouped_captions) - 1:
                transition = transitions[-1]
            else:
                # Select a middle transition based on position
                idx = min(i + 3, len(transitions) - 3)
                transition = transitions[idx]
            
            # Combine similar captions into a coherent sentence
            if len(group) == 1:
                paragraph = f"{transition} {group[0]}"
            else:
                # Extract key elements from similar captions
                combined = group[0]
                for caption in group[1:]:
                    # Find unique elements in this caption
                    current_words = set(combined.lower().split())
                    new_words = set(caption.lower().split())
                    unique_words = new_words - current_words
                    
                    if unique_words:
                        # Add unique elements to the combined description
                        unique_phrase = ' '.join([w for w in caption.split() if w.lower() in unique_words])
                        combined += f" with {unique_phrase}"
                
                paragraph = f"{transition} {combined}"
            
            paragraphs.append(paragraph)
        
        # Combine paragraphs into a flowing narrative
        narrative = ". ".join(paragraphs) + "."
        
        # Clean up any double periods or spacing issues
        narrative = narrative.replace("..", ".").replace(" .", ".").replace(".", ". ").strip()
        
        summary += narrative
        
        # Add audio transcription to the summary if available
        if 'audio_transcription' in locals() and audio_transcription and not audio_transcription.startswith("Error") and not audio_transcription.startswith("Could not"):
            # Process the transcription to make it more readable
            processed_transcription = audio_transcription.strip()
            print("\nAudio Transcription:")
            print(processed_transcription)
            
            # Add the transcription section to the summary
            summary += "\n\nAudio Transcription:\n\n"
            
            # If the transcription is very long, summarize it
            if len(processed_transcription.split()) > 200:
                # Use the existing summarization function to create a concise summary of the transcription
                print("Summarizing audio transcription...")
                try:
                    # Initialize a summarization model
                    model_name = "facebook/bart-large-cnn"
                    summarizer = pipeline("summarization", model=model_name, tokenizer=model_name, framework="pt", device=0 if torch.cuda.is_available() else -1)
                    
                    # Split text into manageable chunks and summarize each chunk
                    text_chunks = chunk_text(processed_transcription)
                    transcription_summaries = []
                    
                    for chunk in text_chunks:
                        if len(chunk.split()) > 10:  # Only summarize if there's enough content
                            summary_output = summarizer(chunk, max_length=100, min_length=30, do_sample=False)
                            transcription_summaries.append(summary_output[0]['summary_text'])
                        else:
                            transcription_summaries.append(chunk)
                    
                    # Combine summaries
                    processed_transcription = " ".join(transcription_summaries)
                except Exception as e:
                    logging.error(f"Transcription summarization error: {e}")
                    # If summarization fails, use the original transcription
                    pass
            
            summary += processed_transcription
            
            # Create a combined summary section that integrates visual and audio information
            summary += "\n\nIntegrated Summary:\n\n"
            
            # Create a more coherent narrative by combining visual and audio information
            try:
                # Initialize a summarization model for the final integration
                model_name = "facebook/bart-large-cnn"
                integrator = pipeline("summarization", model=model_name, tokenizer=model_name, framework="pt", device=0 if torch.cuda.is_available() else -1)
                
                # Combine visual narrative and processed transcription
                combined_text = f"Visual content: {narrative} Audio content: {processed_transcription}"
                
                # Generate an integrated summary
                integrated_summary = integrator(combined_text, max_length=150, min_length=50, do_sample=False)
                summary += integrated_summary[0]['summary_text']
            except Exception as e:
                logging.error(f"Integration summarization error: {e}")
                # If integration fails, provide a simple combined summary
                summary += "This video contains visual elements showing " + narrative.lower() + " The audio discusses " + processed_transcription[:100] + "..."
        
        return summary
        
    except Exception as e:
        logging.error(f"Video summarization error: {e}")
        return f"Error during video summarization: {str(e)}"
    
def answer_question(text, ask):
    """
    Answer questions based on the given text using a pre-trained QA model.
    
    :param text: The original document content.
    :param question: The question asked by the user.
    :return: Answer extracted from the text.
    """
    try:
        qa_model = pipeline("question-answering", model="distilbert-base-cased-distilled-squad")
        result = qa_model(question=ask, context=text)
        return result["answer"]
    except Exception as e:
        return f"Error processing question: {str(e)}"


# Output Functions
def save_as_txt(text, filename="summary_output.txt"):
    with open(filename, "w", encoding="utf-8") as f:
        f.write(text)

def save_as_json(data, filename="summary_output.json"):
    with open(filename, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=4)

def save_as_pdf(text, filename="summary_output.pdf"):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, text)
    pdf.output(filename)

def save_as_docx(text, filename="summary_output.docx"):
    doc = Document()
    doc.add_paragraph(text)
    doc.save(filename)

# Main Application
def main():
    print("Welcome to the Advanced Text Summarizer!")
    print("Choose a summarization model (t5-small/bart-large/distilbart/pegasus/t5-3b):")
    model_choice = input("Model: ")
    user_input = ""

    while True:
        print("\nOptions: text, file, pdf, image, audio, video, exit")
        choice = input("Enter your choice: ").lower()

        if choice == 'exit':
            print("Exiting. Goodbye!")
            break

        elif choice == 'text':
            user_input = input("Enter text: ")
        
        elif choice == 'file':
            filename = input("Enter filename (.txt): ")
            with open(filename, "r", encoding="utf-8") as f:
                user_input = f.read()

        elif choice == 'pdf':
            filename = input("Enter PDF filename: ")
            user_input = read_pdf(filename)

        elif choice == 'image':
            filename = input("Enter image filename: ")
            user_input = extract_text_from_image(filename)
            if "Error" in user_input:
                print(user_input)
                continue


        elif choice == 'audio':
            filename = input("Enter audio filename: ")
            user_input = transcribe_audio(filename)
            
        elif choice == 'video':
            filename = input("Enter video filename: ")
            user_input = summarize_video(filename)
            if "Error" in user_input:
                print(user_input)
                continue
        
        else:
            print("Invalid choice.")
            continue

        if not user_input.strip():
            print("Empty input! Try again.")
            continue

        language = detect(user_input)
        if language != "en":
            print(f"Detected language: {language}, translating to English...")
            user_input = GoogleTranslator(source=language, target="en").translate(user_input)


        processed_text = preprocess_text(user_input)
        summary = summarize_text(processed_text, model_choice)


        print("\nSummary:")
        print(summary)

        save_format = input("Save as (txt/json/pdf/docx): ").lower()
        if save_format == 'txt':
            save_as_txt(summary)
        elif save_format == 'json':
            save_as_json({"summary": summary})
        elif save_format == 'pdf':
            save_as_pdf(summary)
        elif save_format == 'docx':
            save_as_docx(summary)
        else:
            print("Invalid save format.")
        
        while True:
            ask_choice = input("Do you want to ask a question about the text? (yes/no): ").lower()
            if ask_choice == 'yes':
                ask = input("Enter your question: ")
                answer = answer_question(user_input, ask)
                print("\nAnswer:\n", answer)
            elif ask_choice == 'no':
                break
            else:
                print("Invalid input! Please type 'yes' or 'no'.")

if __name__ == "__main__":
    main()
>>>>>>> 332449a7ae3ed30e52c34f172658560db2ded9b0
